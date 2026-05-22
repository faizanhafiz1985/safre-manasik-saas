"""
Generates a comprehensive Word document describing the Safre Manasik SaaS:
architecture, platforms+costs, source code walk-through, and a self-sufficiency
maintenance guide so the owner can operate without Claude.

Run:  python generate_docs.py
Output: Safre_Manasik_Documentation.docx
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Style helpers ────────────────────────────────────────────────────────────

BRAND_GREEN = RGBColor(0x1B, 0x4B, 0x35)
BRAND_GOLD = RGBColor(0xC9, 0xA2, 0x27)
GREY = RGBColor(0x55, 0x55, 0x55)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if color: run.font.color.rgb = color
        else: run.font.color.rgb = BRAND_GREEN
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    return p

def add_code(doc, code, language=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(doc, headers, rows, header_bg='1B4B35'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_bg(cell, header_bg)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    return table

def page_break(doc):
    doc.add_page_break()

# ─── Source code helpers ─────────────────────────────────────────────────────

ROOT = Path(__file__).parent
def read_file(rel_path, max_lines=None):
    try:
        text = (ROOT / rel_path).read_text(encoding='utf-8', errors='ignore')
        if max_lines:
            lines = text.split('\n')
            if len(lines) > max_lines:
                text = '\n'.join(lines[:max_lines]) + f'\n... ({len(lines)-max_lines} more lines — see source file)'
        return text
    except Exception as e:
        return f'[file not found: {rel_path}]'

# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ─── COVER PAGE ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('\n\n\n\nSAFRE MANASIK')
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = BRAND_GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Umrah & Hajj Travel Management SaaS Platform')
r.font.size = Pt(16)
r.italic = True
r.font.color.rgb = BRAND_GOLD

# Logo placeholder note
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run('\n[Insert company logo here — see frontend/public/logo.svg]')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

# Embed actual SVG note + try to embed logo if a PNG version exists
logo_png_candidates = [
    ROOT / 'logo.png',
    ROOT / 'frontend' / 'public' / 'logo.png',
    ROOT / 'safre_manasik_logo.png',
]
embedded = False
for cand in logo_png_candidates:
    if cand.exists():
        try:
            doc.add_picture(str(cand), width=Inches(4.5))
            # Center the just-added picture
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            embedded = True
            break
        except Exception:
            pass
if not embedded:
    # Decorative branded block as visual placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n\n♦  YOUR JOURNEY, OUR RESPONSIBILITY  ♦')
    r.font.size = Pt(14)
    r.font.color.rgb = BRAND_GOLD
    r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nUMRAH  ·  HAJJ  ·  TRAVEL  ·  ZIARAAT  ·  TOURS')
r.font.size = Pt(11)
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Registered in KSA  ·  Commercial Registration No. 7053347410')
r.font.size = Pt(10)
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n\nComplete Platform Documentation')
r.font.size = Pt(18)
r.bold = True
r.font.color.rgb = BRAND_GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Architecture · Source Code · Platforms · Costs · Maintenance Runbook')
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n\n\n\nVersion 2.0  ·  May 2026')
r.font.size = Pt(12)
r.font.color.rgb = GREY

page_break(doc)

# ─── TABLE OF CONTENTS ───────────────────────────────────────────────────────
add_heading(doc, 'Table of Contents', level=1)
toc = [
    '1.  Executive Summary',
    '2.  Application Overview & Capabilities',
    '3.  System Architecture',
    '4.  Technology Stack',
    '5.  Platforms Used & Monthly Costs',
    '6.  Application Workflow (End-to-End)',
    '7.  User Roles & Permissions',
    '8.  Module-by-Module Walkthrough',
    '9.  Multi-Tenancy & Plan System',
    '10. Source Code Walkthrough',
    '11. Deployment & Live URLs',
    '12. Credentials & Access',
    '13. Maintenance Guide',
    '14. Self-Maintenance Without Claude',
    '15. Troubleshooting',
    '16. Future Roadmap',
    '17. Appendix: Full File Inventory',
]
for entry in toc:
    p = doc.add_paragraph(entry)
    p.paragraph_format.left_indent = Inches(0.3)
    for r in p.runs: r.font.size = Pt(11)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Executive Summary', level=1)

add_para(doc,
    'Safre Manasik is a production-ready, multi-tenant Software-as-a-Service '
    'platform built to manage every operational aspect of an Umrah and Hajj '
    'travel agency: bookings, customer management, hotel and transport '
    'allocation, catering, voucher generation, payments, and end-to-end '
    'reporting. The platform is offered as a hosted SaaS that any Umrah '
    'agency can sign up for and operate independently — with strict data '
    'isolation between agencies enforced at the database layer.')

add_heading(doc, 'What you have today', level=2)
for line in [
    ('Live production deployment ', 'on Railway (PaaS), auto-deployed from GitHub on every push.'),
    ('Multi-tenant SaaS ', 'with row-level isolation — each agency sees only its own data.'),
    ('Three configurable plans ', '(Starter / Growth / Enterprise) with per-plan limits and feature flags, editable at runtime by the platform super-admin without code changes.'),
    ('Self-service tenant signup ', '— new agencies can register themselves at /signup.'),
    ('PayPal payment integration ', '(currently in sandbox / stub mode pending live credentials).'),
    ('Branded vouchers ', '(HTML preview + PDF download for tenants on plans that include the feature).'),
    ('Daily schedule & transport reports ', 'with CSV export.'),
    ('Role-based access control ', '(SUPER_ADMIN, ADMIN, AGENT, CUSTOMER).'),
    ('Quota enforcement ', '— hard limits on users and bookings per plan.'),
    ('Feature gating ', '— PDF vouchers, reports, API access, and custom branding can each be turned on or off per plan from the super-admin UI.'),
]:
    add_bullet(doc, line[1], bold_prefix=line[0])

add_heading(doc, 'Where it runs', level=2)
add_para(doc,
    'The entire system runs on Railway, a PaaS provider. The code lives in a '
    'public GitHub repository that auto-deploys to Railway on every push. '
    'PostgreSQL data is hosted on Railway with attached storage. Domains are '
    'registered with Dynadot. PayPal is the payment gateway.')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 2. APPLICATION OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Application Overview & Capabilities', level=1)

add_heading(doc, 'Core operational modules', level=2)
modules = [
    ('Packages', 'Define Umrah / Hajj travel packages with prices, durations, included hotels, and inclusions.'),
    ('Bookings', 'Customers (or staff on behalf of customers) book packages. Track passengers, status, payments, and amenities (transport + catering) assigned per booking.'),
    ('Hotels', 'Maintain a database of partner hotels with rooms, rates, and availability windows.'),
    ('Transport', 'Manage vehicles, routes, and per-booking transport allocation.'),
    ('Catering', 'Manage catering vendors and per-booking meal plan allocation.'),
    ('Vouchers', 'Generate branded HTML and PDF vouchers for confirmed bookings.'),
    ('Payments & Invoices', 'Record payments, generate invoices, integrate with PayPal for online payment.'),
    ('Reports', 'Daily Schedule (all events for a given day) and Transport report (vehicle utilisation over a date range). CSV export.'),
    ('Users', 'Tenant admins create staff accounts (agents, customers).'),
    ('Tenant Settings', 'Each tenant configures its own branding (logo, colours), contact details, currency, and language.'),
    ('Platform Admin', 'SUPER_ADMIN sees all tenants and platform-wide stats; can edit any tenant, suspend / activate, and configure plans.'),
    ('Plans & Pricing', 'SUPER_ADMIN configures the Starter / Growth / Enterprise plans: max users, max bookings, monthly price, and feature flags (PDF vouchers, reports, API access, custom branding, plus any custom flag key).'),
]
for name, desc in modules:
    add_bullet(doc, desc, bold_prefix=f'{name} — ')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. System Architecture', level=1)

add_heading(doc, 'High-level architecture', level=2)
add_code(doc, '''
┌──────────────┐      HTTPS       ┌────────────────────┐   nginx reverse proxy   ┌──────────────────┐
│   Browser    │ ───────────────▶ │   Frontend (React) │ ──────────────────────▶│  Backend (Node)  │
│ (any device) │                  │  served by nginx   │   /api/*  proxied       │   Express API    │
└──────────────┘                  └────────────────────┘                         └──────────────────┘
                                                                                          │
                                                                                          │ Prisma ORM
                                                                                          ▼
                                                                                  ┌──────────────────┐
                                                                                  │   PostgreSQL 17  │
                                                                                  │   (Railway       │
                                                                                  │   managed +      │
                                                                                  │   volume)        │
                                                                                  └──────────────────┘
'''.strip())

add_heading(doc, 'Why this design', level=2)
for line in [
    ('Single domain pattern ', 'all tenants log in at the same URL. Tenancy is enforced via JWT-encoded tenantId and Prisma middleware that auto-filters every query by that id. No risk of cross-tenant data leakage — the database row simply isn\'t returned to the wrong tenant.'),
    ('Frontend → nginx → backend ', 'the frontend container ships its own nginx that proxies /api/* to the backend service. This means the browser only ever talks to the frontend origin, sidestepping CORS pain and making custom-domain SSL easier.'),
    ('Stateless backend ', 'JWT-based auth means the backend has no session state. Restarts are free, and we can scale horizontally later by simply adding replicas.'),
    ('Database-driven configuration ', 'plan limits and feature flags live in the PlanConfig table — no business rules baked into code. The super-admin changes data, the system adapts within seconds.'),
    ('Single repo, two services ', 'frontend and backend share one git repo so a single feature can ship through a single PR. Railway treats each top-level folder as a separate deployable service.'),
]:
    add_bullet(doc, line[1], bold_prefix=line[0])

add_heading(doc, 'Multi-tenancy enforcement (how data is isolated)', level=2)
add_para(doc,
    'Every tenant-owned table (users, bookings, packages, hotels, vehicles, '
    'etc.) carries a tenantId column. The backend uses AsyncLocalStorage to '
    'attach the current request\'s tenantId to a context. A Prisma client '
    'middleware reads that context on every query and injects '
    'where: { tenantId } on reads and data: { tenantId } on writes — '
    'automatically and unconditionally. SUPER_ADMIN bypasses this filtering '
    'by setting isSuperAdmin: true in the context, allowing cross-tenant '
    'visibility for platform operations.')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 4. TECH STACK
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Technology Stack', level=1)

add_heading(doc, 'Backend', level=2)
add_table(doc,
    ['Component', 'Technology', 'Purpose'],
    [
        ['Runtime',          'Node.js 20 (LTS)',                   'JavaScript server runtime'],
        ['Framework',        'Express.js',                          'HTTP routing, middleware'],
        ['ORM',              'Prisma',                              'Type-safe DB access, migrations, tenant middleware'],
        ['Database',         'PostgreSQL 17',                       'Primary data store'],
        ['Auth',             'JWT (jsonwebtoken) + bcryptjs',       'Stateless authentication, password hashing'],
        ['PDF generation',   'Puppeteer (with HTML fallback)',      'Voucher rendering'],
        ['Payments',         '@paypal/checkout-server-sdk',         'PayPal Orders API integration'],
        ['Logger',           'Winston',                             'Structured logging to stdout (picked up by Railway)'],
        ['Security',         'Helmet, express-rate-limit, CORS',    'CSP headers, brute-force protection, cross-origin policy'],
        ['Validation',       'express-validator',                   'Request body validation'],
    ])

add_heading(doc, 'Frontend', level=2)
add_table(doc,
    ['Component', 'Technology', 'Purpose'],
    [
        ['Framework',        'React 18',                            'Single-page application'],
        ['UI library',       'Material UI (MUI) v5',                'Pre-built components, theming'],
        ['Routing',          'react-router-dom v6',                 'Client-side routing'],
        ['HTTP client',      'Axios',                               'API calls with interceptors for JWT + 401 handling'],
        ['Forms',            'react-hook-form',                     'Form state management'],
        ['Charts',           'Chart.js',                            'Dashboard visualisations'],
        ['Notifications',    'react-toastify',                      'Toast popups'],
        ['Build tool',       'Create React App (react-scripts 5)',  'Webpack bundling'],
        ['Server',           'nginx (alpine)',                      'Serves static assets + proxies /api/* to backend'],
    ])

add_heading(doc, 'Infrastructure', level=2)
add_table(doc,
    ['Component', 'Provider', 'Purpose'],
    [
        ['Source control',     'GitHub',                           'Git repo, CI trigger source'],
        ['Hosting (PaaS)',     'Railway',                          'Runs frontend, backend, and database services with auto-deploy from GitHub'],
        ['Database hosting',   'Railway-managed Postgres 17',      'Daily snapshots on paid plans'],
        ['Domain registrar',   'Dynadot',                          'Owns safremanasik.com'],
        ['SSL certificates',   'Railway (Let\'s Encrypt)',         'Auto-issued + renewed for custom domains'],
        ['Payment gateway',    'PayPal',                           'Sandbox today, live when credentials are added'],
        ['Currency conversion','In-app (SAR↔USD pegged at 3.75)',  'PayPal does not support SAR; backend auto-converts'],
    ])

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 5. PLATFORMS + COSTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Platforms Used & Monthly Costs', level=1)

add_para(doc,
    'Below is the complete platform / vendor list with what each one does and '
    'what it costs you. Prices are USD unless stated otherwise and reflect '
    'public pricing as of May 2026 — always confirm current tiers on each '
    'provider\'s site.', italic=True)

add_table(doc,
    ['Platform', 'Purpose', 'Plan tier needed', 'Approx. cost / month', 'Mandatory?'],
    [
        ['GitHub',            'Source code hosting & CI trigger',                  'Free (public repo)',                    'USD 0',                       'Yes'],
        ['Railway',           'Hosts frontend, backend, and PostgreSQL',           'Hobby ($5/mo) or Pro ($20/mo)',         'USD 5–20 + usage',            'Yes'],
        ['Railway Postgres',  'Managed PostgreSQL with attached volume',           'Included in Railway plan above',        '~USD 5 storage + traffic',    'Yes'],
        ['Dynadot',           'Domain registrar for safremanasik.com',             'Annual (~USD 10.88 / year)',            'USD ~1',                      'Yes (for custom domain)'],
        ['PayPal',            'Customer payments (cards, PayPal balance, etc.)',   'Free account, per-transaction fees',    'USD 0 fixed + ~3.5% per txn', 'Optional (until you accept payments)'],
        ['Cloudflare (option)','Alternative DNS — fixes Dynadot TXT issue',        'Free tier',                             'USD 0',                       'Optional (recommended)'],
        ['Postmark / Resend (option)', 'Outgoing transactional email',             'Postmark $15/mo or Resend free tier',   'USD 0–15',                    'Optional (recommended later)'],
        ['Sentry / Logflare (option)', 'Error tracking + log aggregation',         'Free / pay-as-you-go',                  'USD 0–10',                    'Optional'],
    ])

add_heading(doc, 'Estimated monthly run cost', level=2)
add_para(doc,
    'For a starting deployment with 1-5 tenants and modest usage you should '
    'budget approximately:')
add_table(doc,
    ['Scenario', 'Monthly cost (USD)', 'Notes'],
    [
        ['Bare minimum (trial only)',           '~$5',     'Railway Hobby + free domain + no PayPal yet'],
        ['Recommended live setup',              '~$25',    'Railway Hobby + custom domain + Postmark email'],
        ['Scaling to 50+ tenants',              '~$50–100', 'Railway Pro plan + transactional email + paid Sentry'],
        ['PayPal',                              '0 + ~3.5% per txn', 'Pay-as-you-go, no monthly fee'],
    ])

add_heading(doc, 'Annual fixed costs', level=2)
add_bullet(doc, 'Domain (safremanasik.com on Dynadot): ~USD 11 / year')
add_bullet(doc, 'Optional: SSL certificate (free via Railway / Cloudflare / Let\'s Encrypt) — USD 0')
add_bullet(doc, 'Optional: Custom transactional email domain — depends on provider')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 6. WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Application Workflow (End-to-End)', level=1)

add_heading(doc, '6.1 New tenant sign-up flow', level=2)
add_code(doc, '''
Customer visits  →  /signup page (single URL, no subdomain)
       │
       ▼
Fills form: organisation name, admin email, password
       │
       ▼
POST /api/auth/signup-tenant
       │
       ▼
Backend creates:  Tenant row (slug auto-generated)
                  User row (role=ADMIN, tenantId=<new id>)
                  Returns JWT
       │
       ▼
Frontend stores JWT in localStorage  →  redirects to /dashboard
       │
       ▼
New tenant lands in empty workspace ready to add packages, bookings, etc.
'''.strip())

add_heading(doc, '6.2 Daily operations flow (typical tenant agent)', level=2)
add_code(doc, '''
Agent logs in  →  /dashboard shows tenant-scoped stats
   │
   ├─▶  Packages page  →  Define / edit Umrah packages
   │
   ├─▶  Bookings page  →  Create booking for a customer
   │                      • Pick package
   │                      • Add passengers
   │                      • Assign transport (vehicle + route)
   │                      • Assign catering (vendor + meal plan)
   │                      • Generate voucher (HTML always; PDF if plan includes pdfVouchers)
   │                      • Record payments (or send PayPal link to customer)
   │
   ├─▶  Reports / Daily Schedule  →  see today's check-ins, check-outs, transports
   │                                  (gated behind 'reports' feature flag)
   │
   └─▶  Profile  →  change own password
'''.strip())

add_heading(doc, '6.3 Payment flow (customer pays via PayPal)', level=2)
add_code(doc, '''
Booking detail page  →  click "Pay with PayPal"
   │
   ▼
Frontend: POST /api/payments/gateway/paypal/create-order
   │
   ▼
Backend: converts SAR amount → USD at 3.75 peg
   │     → calls PayPal Orders API
   │     → returns approval URL
   ▼
Browser redirected to PayPal
   │
   ▼
Customer approves on PayPal
   │
   ▼
PayPal redirects to /payment/paypal/success?token=ORDER_ID
   │
   ▼
Frontend: POST /api/payments/gateway/paypal/capture-order
   │
   ▼
Backend: captures, records Payment row, updates Invoice.status (PARTIAL/PAID)
   │
   ▼
"Payment Successful" page → redirect back to booking
'''.strip())

add_heading(doc, '6.4 SUPER_ADMIN platform-admin flow', level=2)
add_code(doc, '''
superadmin@safremanasik.com logs in  →  redirected to /super-admin
   │
   ├─▶  Platform stats:  total tenants, users, bookings, revenue
   │
   ├─▶  Tenant list:    edit any tenant, suspend, activate, delete
   │                     edit per-tenant maxUsers / maxBookings overrides
   │
   └─▶  Plans & Pricing:  3 plan cards (Starter / Growth / Enterprise)
                          • Edit display name, description, default limits
                          • Toggle feature flags (PDF, reports, API, branding)
                          • Add new custom feature keys on the fly
                          • Set monthly price + currency
                          → Save → applies platform-wide within 5 seconds
'''.strip())

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 7. ROLES
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. User Roles & Permissions', level=1)

add_table(doc,
    ['Role', 'Scope', 'Typical actions', 'Restrictions'],
    [
        ['SUPER_ADMIN', 'Platform-wide (no tenant)', 'Manage all tenants, configure plans, view platform stats', 'Cannot belong to a tenant; intentionally has no booking workflow'],
        ['ADMIN',       'Their tenant only',           'Manage packages, hotels, users, settings, bookings, vouchers, payments, reports', 'Cannot see other tenants\' data'],
        ['AGENT',       'Their tenant only',           'Create / edit bookings, vouchers, payments, transport, catering', 'Cannot manage users, tenant settings, or delete bookings'],
        ['CUSTOMER',    'Their tenant only',           'View own bookings, vouchers, packages',                          'No create/edit on operational entities'],
    ])

add_para(doc,
    'Role enforcement happens in two places: the JWT carries the role claim, '
    'and Express middleware (authorize(\'ADMIN\', \'AGENT\')) gates each '
    'route. Combined with tenant isolation, this means a malicious agent '
    'cannot impersonate another tenant\'s admin even with a forged request — '
    'they\'d still be blocked on the role check and the tenant filter.', italic=True)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 8. MODULES
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Module-by-Module Walkthrough', level=1)

module_details = [
    ('Auth (backend/src/routes/auth.js + authController.js)',
     'JWT-based authentication. Endpoints: signup-tenant (creates tenant + first admin), register (creates user within tenant), login, me, change-password, update profile. Passwords hashed with bcrypt (12 rounds). Rate-limited (30 attempts / 15 min) to slow brute force.'),
    ('Users (backend/src/routes/users.js + userController.js)',
     'Tenant admins create staff accounts. Quota-enforced — when a tenant hits its plan\'s maxUsers limit, the create endpoint returns 403 with an upgrade message. Roles allowed: ADMIN, AGENT, CUSTOMER.'),
    ('Packages (backend/src/routes/packages.js + packageController.js)',
     'CRUD for Umrah / Hajj packages: name, description, duration, departure city, departure date, prices in tenant\'s currency. Supports nested package_hotels and price_tiers.'),
    ('Bookings (backend/src/routes/bookings.js + bookingController.js)',
     'Core operational entity. Quota-enforced. Supports nested passenger list, transport assignment, catering assignment. Status transitions (PENDING → CONFIRMED → COMPLETED / CANCELLED).'),
    ('Hotels (backend/src/routes/hotels.js + hotelController.js)',
     'Hotel database with rooms, rates, city, country. Linked to packages via PackageHotel.'),
    ('Transport (backend/src/routes/transport.js + transportController.js)',
     'Vehicles + routes. Booking-transport assignment records which vehicle takes which booking, when.'),
    ('Catering (backend/src/routes/catering.js + cateringController.js)',
     'Catering vendors and meal plans. Booking-catering assignment.'),
    ('Vouchers (backend/src/routes/vouchers.js + voucherController.js)',
     'Generates branded HTML voucher for any booking. PDF generation gated behind the pdfVouchers feature flag. PDF rendering uses Puppeteer; falls back to plain HTML if Chromium can\'t launch.'),
    ('Payments (backend/src/routes/payments.js + paymentController.js + paymentGateway.js)',
     'Manual payment recording + PayPal Orders API integration. SAR is converted to USD at a 3.75 peg for PayPal calls (PayPal doesn\'t accept SAR). Webhook endpoint captures asynchronous payment events.'),
    ('Reports (backend/src/routes/reports.js + reportController.js)',
     'Daily Schedule (events for a date) and Transport Report (vehicle utilisation over a range). CSV export. Gated behind reports feature flag.'),
    ('Tenant (backend/src/routes/tenant.js + tenantController.js)',
     'Tenant self-management: read own tenant info, update settings (name, branding, contact), read current plan + usage (/current/quota).'),
    ('Super Admin (backend/src/routes/superAdmin.js + superAdminController.js)',
     'Platform-admin endpoints: list / get / update / suspend / activate / delete tenants; platform stats; cross-tenant booking list; plan management; per-tenant usage snapshot.'),
    ('Quota middleware (backend/src/middleware/quota.js)',
     'Two reusable middlewares: checkQuota(\'users\'|\'bookings\') blocks creation when a tenant hits its limit; requireFeature(\'flagKey\') blocks any route if the tenant\'s plan doesn\'t include the flag. 5-second in-memory cache so SUPER_ADMIN edits propagate within seconds.'),
    ('Bootstrap (backend/src/bootstrap.js)',
     'Runs on every server boot. Idempotently ensures (1) a SUPER_ADMIN user exists with the password from SUPERADMIN_PASSWORD env var — and resyncs it on each boot, so the env var is the source of truth; (2) default PlanConfig rows exist for STARTER / GROWTH / ENTERPRISE.'),
]
for name, desc in module_details:
    add_bullet(doc, desc, bold_prefix=f'{name}: ')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 9. MULTI-TENANCY & PLANS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Multi-Tenancy & Plan System', level=1)

add_heading(doc, '9.1 Tenant isolation', level=2)
add_para(doc,
    'Every tenant-owned model in backend/prisma/schema.prisma carries a '
    'tenantId column with a foreign key to Tenant. Prisma middleware in '
    'backend/src/config/database.js calls AsyncLocalStorage to read the '
    'current request\'s tenantId from the request context, then automatically '
    'injects where: { tenantId } into every find / count / update / delete '
    'and data: { tenantId } into every create. The only way to bypass is to '
    'set isSuperAdmin: true on the context (which only SUPER_ADMIN endpoints '
    'do).')

add_heading(doc, '9.2 Plan configuration (no code, all data)', level=2)
add_para(doc,
    'PlanConfig is a database table — one row per plan key (STARTER, GROWTH, '
    'ENTERPRISE). Each row has: displayName, description, defaultMaxUsers, '
    'defaultMaxBookings, features (JSON map of flag → boolean), priceMonthly, '
    'priceCurrency, isActive. The SUPER_ADMIN edits this from the Plans & '
    'Pricing page; changes apply platform-wide within ~5 seconds (cache TTL).')

add_heading(doc, '9.3 Quota enforcement', level=2)
add_para(doc, 'Wired into these routes today:')
add_bullet(doc, 'POST /api/users  → checkQuota(\'users\')')
add_bullet(doc, 'POST /api/bookings → checkQuota(\'bookings\')')
add_bullet(doc, 'POST /api/vouchers/generate → requireFeature(\'pdfVouchers\')')
add_bullet(doc, 'GET  /api/vouchers/download/:id → requireFeature(\'pdfVouchers\')')
add_bullet(doc, 'All /api/reports/* → requireFeature(\'reports\')')

add_para(doc,
    'To gate a new route behind a new feature flag, the only code change is '
    'to add requireFeature(\'yourNewKey\') to the route. Then in the Plans UI '
    'add the custom key and toggle it on for the plans that should have it.')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 10. SOURCE CODE WALKTHROUGH
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '10. Source Code Walkthrough', level=1)

add_para(doc,
    'This section shows the most important source files with a brief purpose '
    'for each. The full source lives at:', italic=True)
add_para(doc, 'https://github.com/faizanhafiz1985/safre-manasik-saas', bold=True, color=BRAND_GREEN)

# Server entry point
add_heading(doc, '10.1 backend/src/server.js — Express entry point', level=2)
add_para(doc, 'Sets up middleware (helmet, cors, compression, rate limiting), wires up all /api/* routes, starts listening, runs the startup bootstrap.', italic=True)
add_code(doc, read_file('backend/src/server.js', max_lines=90))

add_heading(doc, '10.2 backend/prisma/schema.prisma — Database schema', level=2)
add_para(doc, 'Defines every model. Notable: enums TenantPlan / TenantStatus, the PlanConfig table for configurable plans, and tenantId on every tenant-owned model.', italic=True)
add_code(doc, read_file('backend/prisma/schema.prisma', max_lines=140))

add_heading(doc, '10.3 backend/src/config/tenantContext.js — Tenant-scoped async context', level=2)
add_para(doc, 'Uses AsyncLocalStorage to make req.user.tenantId available to deep Prisma middleware without threading it through every function.', italic=True)
add_code(doc, read_file('backend/src/config/tenantContext.js'))

add_heading(doc, '10.4 backend/src/middleware/quota.js — Quota + feature enforcement', level=2)
add_para(doc, 'The two middlewares that gate routes by plan: checkQuota and requireFeature.', italic=True)
add_code(doc, read_file('backend/src/middleware/quota.js'))

add_heading(doc, '10.5 backend/src/bootstrap.js — Startup bootstrap', level=2)
add_para(doc, 'Self-healing creation of the SUPER_ADMIN user and the default plan configs on every boot.', italic=True)
add_code(doc, read_file('backend/src/bootstrap.js'))

add_heading(doc, '10.6 backend/src/controllers/superAdminController.js — Platform admin', level=2)
add_para(doc, 'All SUPER_ADMIN endpoints in one file. Tenant CRUD, suspend / activate, platform stats, plan CRUD, per-tenant usage.', italic=True)
add_code(doc, read_file('backend/src/controllers/superAdminController.js', max_lines=120))

add_heading(doc, '10.7 backend/src/routes/superAdmin.js — Route wiring', level=2)
add_code(doc, read_file('backend/src/routes/superAdmin.js'))

add_heading(doc, '10.8 backend/src/controllers/authController.js — Auth flows', level=2)
add_para(doc, 'Login, signup-tenant, register, change-password, profile. Bcrypt + JWT.', italic=True)
add_code(doc, read_file('backend/src/controllers/authController.js', max_lines=80))

add_heading(doc, '10.9 frontend/src/services/api.js — Axios client', level=2)
add_para(doc, 'Defines the base URL (same-origin /api by default — nginx proxies to backend) and attaches the JWT to every outgoing request.', italic=True)
add_code(doc, read_file('frontend/src/services/api.js'))

add_heading(doc, '10.10 frontend/src/App.js — React router root', level=2)
add_code(doc, read_file('frontend/src/App.js', max_lines=90))

add_heading(doc, '10.11 frontend/src/pages/SuperAdminPlansPage.js — Plan management UI', level=2)
add_para(doc, 'The page where SUPER_ADMIN edits plan limits and feature flags.', italic=True)
add_code(doc, read_file('frontend/src/pages/SuperAdminPlansPage.js', max_lines=100))

add_heading(doc, '10.12 frontend/nginx.conf — Frontend nginx config', level=2)
add_para(doc, 'SPA fallback (every route serves index.html) and the /api proxy to the backend service.', italic=True)
add_code(doc, read_file('frontend/nginx.conf'))

add_heading(doc, '10.13 backend/Dockerfile — Backend container', level=2)
add_code(doc, read_file('backend/Dockerfile'))

add_heading(doc, '10.14 frontend/Dockerfile — Frontend container', level=2)
add_code(doc, read_file('frontend/Dockerfile'))

add_heading(doc, '10.15 docker-compose.yml — Local dev orchestration', level=2)
add_code(doc, read_file('docker-compose.yml'))

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 11. DEPLOYMENT & LIVE URLS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '11. Deployment & Live URLs', level=1)

add_heading(doc, 'Production URLs (today)', level=2)
add_table(doc,
    ['Service', 'URL', 'Status'],
    [
        ['Frontend (live app)', 'https://frontend-production-56ba6.up.railway.app', 'ONLINE'],
        ['Backend (API)',       'https://backend-production-44fd.up.railway.app',  'ONLINE'],
        ['Custom frontend',     'https://app.safremanasik.com',                     'Pending Dynadot TXT record (use workaround in §15)'],
        ['Custom API',          'https://api.safremanasik.com',                     'Pending TXT (currently SSL pending)'],
        ['GitHub repo',         'https://github.com/faizanhafiz1985/safre-manasik-saas', 'Public'],
        ['Railway project',     'zesty-elegance (ID: adcd3710-a113-457d-af99-8a462aa04fd6)', 'Active'],
    ])

add_heading(doc, 'Deployment pipeline', level=2)
add_code(doc, '''
git push origin main   on local machine
        │
        ▼
GitHub receives push, fires webhook to Railway
        │
        ▼
Railway pulls latest commit, builds Docker image for each service
   ├── frontend  : multi-stage build (Node alpine → nginx alpine)
   └── backend   : Node alpine + prisma db push on container start
        │
        ▼
Health check runs on new container
        │
        ▼
Traffic switched to new container, old one drained
        │
        ▼
Live in production (~3 minutes end-to-end)
'''.strip())

add_para(doc, 'NOTE: Railway\'s GitHub auto-deploy has been intermittent. If a push does not trigger a deploy within 5 minutes, change any env variable (it can be a no-op like setting DEPLOY_NONCE to the current timestamp) which forces a rebuild from the latest commit. Or use the "Redeploy" option in the Railway service dashboard.', italic=True)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 12. CREDENTIALS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '12. Credentials & Access', level=1)

add_para(doc,
    'IMPORTANT — keep this section locked. The values here grant full access '
    'to your platform. The same information lives in CREDENTIALS_PROD.txt at '
    'the repo root which is git-ignored. Do not commit that file or share '
    'it.', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

add_heading(doc, 'Production login accounts', level=2)
add_table(doc,
    ['Role', 'Email', 'Password (treat as secret)', 'Login URL'],
    [
        ['SUPER_ADMIN',  'superadmin@safremanasik.com', 'SuperAdmin@LwPSLqHQE65GnZYo!', 'https://frontend-production-56ba6.up.railway.app/login'],
        ['Tenant admin (Safre Manasik)', 'admin@safremanasik.com', 'BY3dFB3xd8zRwPMwA1!', 'same'],
    ])

add_heading(doc, 'Important environment variables (backend service)', level=2)
add_table(doc,
    ['Variable', 'Purpose', 'Where to change'],
    [
        ['DATABASE_URL',         'Postgres connection string',                        'Auto-set by Railway when Postgres service is linked'],
        ['JWT_SECRET',           'Signs JWT tokens — 48+ random bytes',               'Railway → backend → Variables'],
        ['JWT_EXPIRES_IN',       'Token lifetime (e.g. 7d)',                          'Same'],
        ['NODE_ENV',             'production',                                        'Same'],
        ['FRONTEND_URL',         'CORS allowed origin',                               'Same'],
        ['PAYPAL_MODE',          'sandbox | live',                                    'Same'],
        ['PAYPAL_CLIENT_ID',     'PayPal API key',                                    'Same — get from developer.paypal.com'],
        ['PAYPAL_CLIENT_SECRET', 'PayPal API secret',                                 'Same'],
        ['SUPERADMIN_EMAIL',     'Super admin login email (default: superadmin@safremanasik.com)', 'Same'],
        ['SUPERADMIN_PASSWORD',  'Super admin password — synced on every boot. Change this to rotate the super admin password.', 'Same'],
    ])

add_heading(doc, 'Frontend env vars', level=2)
add_table(doc,
    ['Variable', 'Purpose'],
    [
        ['PORT',               'nginx listen port (set by Railway)'],
        ['BACKEND_URL',        'nginx upstream for /api proxy — should point to the backend service (e.g. https://backend-production-44fd.up.railway.app)'],
        ['REACT_APP_API_URL',  'API base URL baked into the React bundle at build time. Set to /api for same-origin (recommended).'],
    ])

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 13. MAINTENANCE GUIDE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '13. Maintenance Guide', level=1)

add_heading(doc, '13.1 Routine tasks (do these regularly)', level=2)
add_table(doc,
    ['Frequency', 'Task', 'How'],
    [
        ['Daily',     'Eyeball Railway services',                'railway.com dashboard → all three services should be Online (green dot)'],
        ['Weekly',    'Check error logs',                         'Railway → backend service → Logs tab. Look for ERROR or 500-class events.'],
        ['Weekly',    'Verify backups',                          'Railway → Postgres service → Backups tab. Confirm daily snapshot exists.'],
        ['Monthly',   'Rotate JWT_SECRET',                        'Railway env var → generate a new 48-byte value. Existing JWTs will be invalidated (users must re-login).'],
        ['Monthly',   'Review tenant list',                      'Login as SUPER_ADMIN → /super-admin → suspend / activate any inactive tenants.'],
        ['Quarterly', 'Update Node.js + npm deps',               '`cd backend && npm outdated` — patch security advisories first.'],
        ['Annually',  'Renew Dynadot domain',                    'Dynadot dashboard → My Domains → renew. Auto-renew available.'],
    ])

add_heading(doc, '13.2 Common maintenance scenarios', level=2)

add_para(doc, 'A. To change a plan\'s limits or features at runtime', bold=True)
add_bullet(doc, 'Login as SUPER_ADMIN → sidebar → Plans & Pricing')
add_bullet(doc, 'Click the edit (pencil) icon on the plan card')
add_bullet(doc, 'Adjust users, bookings, feature toggles, or pricing → Save')
add_bullet(doc, 'Changes apply within 5 seconds platform-wide. No redeploy.')

add_para(doc, 'B. To override limits for one specific tenant only', bold=True)
add_bullet(doc, 'Login as SUPER_ADMIN → /super-admin (Platform Admin)')
add_bullet(doc, 'Find the tenant row → click pencil')
add_bullet(doc, 'Set Max Users / Max Bookings to the override — Save')
add_bullet(doc, 'Tenant-level override beats the plan default')

add_para(doc, 'C. To add a brand-new tenant (manual, by you)', bold=True)
add_code(doc, '''
curl -X POST https://backend-production-44fd.up.railway.app/api/auth/signup-tenant \\
  -H "Content-Type: application/json" \\
  -d '{
    "tenantName": "Sample Travels",
    "tenantSlug": "sample-travels",
    "adminName": "Ali Sample",
    "adminEmail": "admin@sample.com",
    "adminPassword": "Strong@2026!"
  }'

# Or: send your customer the public signup link
# https://frontend-production-56ba6.up.railway.app/signup
'''.strip())

add_para(doc, 'D. To switch PayPal from sandbox to live', bold=True)
add_bullet(doc, 'Get live Client ID + Secret from https://developer.paypal.com/dashboard/applications/live')
add_bullet(doc, 'Railway → backend service → Variables → Update PAYPAL_MODE=live, PAYPAL_CLIENT_ID, PAYPAL_CLIENT_SECRET')
add_bullet(doc, 'Click "Deploy Changes" to apply')
add_bullet(doc, 'Test by completing a 1 SAR sandbox booking first')

add_para(doc, 'E. To rotate the SUPER_ADMIN password', bold=True)
add_bullet(doc, 'Railway → backend service → Variables → Edit SUPERADMIN_PASSWORD to your new value')
add_bullet(doc, 'Save / Deploy Changes')
add_bullet(doc, 'On next boot, bootstrap.js automatically resyncs the DB user\'s password from the env var')
add_bullet(doc, 'Login with the new password')

add_para(doc, 'F. To add a new feature flag for plan gating', bold=True)
add_bullet(doc, 'Login as SUPER_ADMIN → Plans & Pricing')
add_bullet(doc, 'Edit any plan → in the "Add custom feature key" input type the new key (e.g. whatsAppIntegration) → Add → toggle on for plans that should have it')
add_bullet(doc, 'In code, add requireFeature(\'whatsAppIntegration\') to the route(s) that should be gated → commit → push')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 14. SELF-MAINTENANCE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '14. Self-Maintenance Without Claude', level=1)

add_para(doc,
    'This section is the most important. Every operation below is something '
    'you can do yourself — through dashboards or a single command — without '
    'any AI assistance.', italic=True)

add_heading(doc, '14.1 The four places you need to know', level=2)
add_table(doc,
    ['Place', 'URL', 'When you go there'],
    [
        ['Railway dashboard',  'railway.com',                  'Restart services, view logs, change env vars, see metrics, run the redeploy / rollback'],
        ['GitHub repo',        'github.com/faizanhafiz1985/safre-manasik-saas', 'Source code, history, raise issues, edit files via the GitHub web UI'],
        ['Dynadot dashboard',  'dynadot.com',                  'Renew domain, edit DNS records'],
        ['PayPal Developer',   'developer.paypal.com',         'Get live credentials, view transactions'],
    ])

add_heading(doc, '14.2 If the live app goes down', level=2)
add_bullet(doc, 'Open Railway → check status of all three services. Any not green?')
add_bullet(doc, 'If a service is "Crashed": open Logs tab → read the last 50 lines. Common causes:')
add_bullet(doc, 'DATABASE_URL changed → relink Postgres service in Architecture view', level=1)
add_bullet(doc, 'Out of memory → upgrade Railway plan or shrink build', level=1)
add_bullet(doc, 'Bad commit → click kebab on the previous successful deploy → "Redeploy"', level=1)
add_bullet(doc, 'If just slow: Restart the service from the kebab menu')
add_bullet(doc, 'If totally lost: rollback to a known-good deployment from the History panel')

add_heading(doc, '14.3 If you forget the super-admin password', level=2)
add_bullet(doc, 'Railway → backend → Variables → SUPERADMIN_PASSWORD → set to a new strong value')
add_bullet(doc, 'Click "Deploy Changes" — backend restarts')
add_bullet(doc, 'Bootstrap detects the existing SUPER_ADMIN row, resyncs the password to the new value, log in with it')
add_bullet(doc, 'This works because bootstrap.js was written to be self-healing (see §10.5)')

add_heading(doc, '14.4 If you need to roll back to a previous version', level=2)
add_bullet(doc, 'Railway → service → Deployments tab → scroll the History')
add_bullet(doc, 'Find the deployment you want to restore → click its kebab menu → Redeploy')
add_bullet(doc, 'It will replace the active deployment within ~30 seconds')

add_heading(doc, '14.5 If you need to view / export data', level=2)
add_bullet(doc, 'Quickest: login as SUPER_ADMIN → /super-admin → All Tenants or /super-admin/bookings')
add_bullet(doc, 'Direct DB access: Railway → Postgres service → Connect tab — copy the public URL and use any Postgres client (pgAdmin, TablePlus, DBeaver)')
add_bullet(doc, 'For bookings CSV: any tenant ADMIN can hit /reports/daily-schedule/export?date=YYYY-MM-DD')

add_heading(doc, '14.6 If you need to do a code change yourself', level=2)
add_para(doc,
    'You don\'t need to be a developer. Most everyday changes are config '
    'changes — done in the Plans & Pricing page or Tenant Settings. If you '
    'truly must edit code:')
add_bullet(doc, 'Go to github.com/faizanhafiz1985/safre-manasik-saas')
add_bullet(doc, 'Find the file → click the pencil to edit in the browser')
add_bullet(doc, 'Make the change → scroll down → "Commit changes"')
add_bullet(doc, 'Railway sees the new commit and rebuilds automatically (sometimes you must trigger a "Redeploy" manually)')
add_para(doc,
    'Caveat: for any non-trivial change, ask another developer or a freelance '
    'Node.js / React engineer to review before you push to main. The repo is '
    'standard Node.js + React — any competent contractor can pick it up.', italic=True)

add_heading(doc, '14.7 Finding a developer to maintain it', level=2)
add_bullet(doc, 'Skills required: Node.js, Express, Prisma, React, basic Postgres. All common technologies.')
add_bullet(doc, 'Where to look: Upwork, Toptal, local KSA tech communities, Stack Overflow Jobs')
add_bullet(doc, 'Typical maintainer cost: USD 20–60 / hour. Most maintenance asks (config tweaks, new fields, fixing a button) take 1–4 hours.')
add_bullet(doc, 'Give them this document + the GitHub repo + a Railway invite — they have everything they need.')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 15. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '15. Troubleshooting', level=1)

add_table(doc,
    ['Symptom', 'Likely cause', 'Fix'],
    [
        ['"An error occurred" on login',
         'Frontend can\'t reach the backend (SSL or proxy issue)',
         'Confirm REACT_APP_API_URL=/api and BACKEND_URL points to a working backend URL. Hard-refresh browser.'],
        ['502 Bad Gateway from frontend',
         'Backend down, or nginx BACKEND_URL points to a domain without a valid SSL cert',
         'Check backend service status in Railway. If using custom api.subdomain that lacks SSL, point BACKEND_URL to the .up.railway.app URL instead.'],
        ['"Invalid credentials" for SUPER_ADMIN',
         'Database row out of sync with env var',
         'Set SUPERADMIN_PASSWORD env var to your desired value → redeploy → bootstrap resyncs.'],
        ['"Feature \'X\' is not available on your STARTER plan"',
         'Working as designed — your plan doesn\'t include that feature',
         'Either upgrade the plan (in the tenant\'s plan column), or as SUPER_ADMIN flip the feature flag in Plans & Pricing.'],
        ['"users limit reached" when creating a user',
         'Tenant has hit maxUsers',
         'SUPER_ADMIN → Edit Tenant → bump Max Users. Or upgrade tenant\'s plan.'],
        ['Custom domain (app.safremanasik.com) returns 404',
         'CNAME points to Railway but the TXT verification record is missing',
         'Dynadot doesn\'t allow the underscore subdomain Railway requires for TXT. Workaround: move DNS to Cloudflare (free).'],
        ['New code committed but Railway didn\'t deploy',
         'GitHub auto-deploy webhook intermittent',
         'Add or change any env var → "Deploy Changes". Or click Redeploy on the active deployment from the kebab menu.'],
        ['PayPal payment fails',
         'Sandbox account out of credit, or live not configured',
         'Sandbox: top up the sandbox buyer at developer.paypal.com → Sandbox accounts. Live: confirm PAYPAL_CLIENT_ID/SECRET match the live app, not sandbox.'],
        ['PDF voucher endpoint returns 403 for tenant',
         'Plan does not include pdfVouchers',
         'Toggle pdfVouchers on for that plan in Plans & Pricing.'],
        ['Logs show "Bootstrap failed: relation does not exist"',
         'Prisma db push didn\'t run',
         'Inspect deploy logs for prisma errors; manually run `npx prisma db push` from Railway shell if needed.'],
    ])

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 16. ROADMAP
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '16. Future Roadmap', level=1)

add_para(doc, 'Suggested improvements, in rough priority order:', italic=True)

add_heading(doc, 'Quick wins (1–4 hours each)', level=2)
add_bullet(doc, 'Outgoing email — wire up Postmark / Resend; send booking confirmations, payment receipts, voucher links')
add_bullet(doc, 'Tenant signup welcome email')
add_bullet(doc, 'Better dashboard charts (occupancy rate, revenue trend)')
add_bullet(doc, 'Activate Cloudflare DNS to unblock the custom domain TXT verification')
add_bullet(doc, 'Switch PayPal to live mode (once business validation is complete)')

add_heading(doc, 'Medium-term (1–3 days each)', level=2)
add_bullet(doc, 'Arabic UI (i18n) — full RTL support')
add_bullet(doc, 'ZATCA e-invoicing for Saudi tax compliance')
add_bullet(doc, 'WhatsApp Business integration for booking notifications')
add_bullet(doc, 'More report types: revenue by package, agent performance, customer LTV')
add_bullet(doc, 'Audit log viewer in the SUPER_ADMIN UI (the audit_logs table is already populated)')

add_heading(doc, 'Larger initiatives (1+ weeks)', level=2)
add_bullet(doc, 'Customer-facing self-service portal (currently customers exist as a role but the UI is minimal)')
add_bullet(doc, 'Mobile app (React Native, sharing the existing API)')
add_bullet(doc, 'Two-factor authentication for ADMIN and SUPER_ADMIN logins')
add_bullet(doc, 'White-label sub-domains per tenant (currently single URL with login-based scoping)')
add_bullet(doc, 'Stripe alongside PayPal for card-only customers')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 17. APPENDIX: FILE INVENTORY
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, '17. Appendix: Full File Inventory', level=1)

add_para(doc,
    'Generated automatically from the repo. Lists every source-code file '
    'with its size. Use it as a map when learning the codebase.', italic=True)

import subprocess
try:
    out = subprocess.check_output(
        ['git', 'ls-files'], cwd=str(ROOT), encoding='utf-8', errors='ignore'
    )
    files = [f for f in out.split('\n') if f and not f.startswith(('node_modules', 'frontend/build', 'backend/logs'))]
    files = [f for f in files if not any(f.endswith(ext) for ext in ['.lock', '.lock.json', '.gitignore'])]
    files.sort()
    rows = []
    for f in files:
        path = ROOT / f
        try:
            size = path.stat().st_size
            sz_str = f'{size:,} B' if size < 1024 else f'{size/1024:.1f} KB'
        except Exception:
            sz_str = '—'
        rows.append([f, sz_str])
    # Word tables get heavy at large sizes — paginate at 200 rows per table
    chunk = 0
    while rows:
        page = rows[:200]
        rows = rows[200:]
        add_table(doc, ['File', 'Size'], page)
        if rows:
            add_para(doc, f'(continued on next page — {len(rows)} more files)', italic=True, color=GREY)
            page_break(doc)
            chunk += 1
except Exception as e:
    add_para(doc, f'[file inventory unavailable: {e}]', italic=True, color=GREY)

# ─── Final footer ────────────────────────────────────────────────────────────
page_break(doc)
add_heading(doc, 'Document end', level=1)
add_para(doc,
    'This document was generated programmatically from the live source tree on ' +
    __import__('datetime').date.today().isoformat() +
    '. To regenerate, run python generate_docs.py from the repo root.',
    italic=True, color=GREY)
add_para(doc,
    'For questions about anything in here: the source code is the ultimate '
    'source of truth — every claim in this document can be verified against '
    'the file paths cited.',
    italic=True, color=GREY)

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = ROOT / 'Safre_Manasik_Documentation.docx'
doc.save(str(output_path))
print(f'OK -> {output_path}')
print(f'   size: {output_path.stat().st_size / 1024:.1f} KB')
